#!/usr/bin/env python3
"""
Quiz Export Module - Generate Word and PDF documents from quiz data
"""

import io
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import black, blue, green, red
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

logger = logging.getLogger(__name__)

@dataclass
class QuizExportData:
    """Data structure for quiz export."""
    title: str
    questions: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    include_answers: bool = True


def create_quiz_word_document(quiz_data: QuizExportData) -> Optional[io.BytesIO]:
    """
    Create a Word document from quiz data.
    
    Args:
        quiz_data: QuizExportData containing quiz information
        
    Returns:
        BytesIO buffer containing the Word document or None if failed
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available for Word document generation")
        return None
    
    try:
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(quiz_data.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run("Generated: ").bold = True
        metadata_para.add_run(f"{quiz_data.metadata.get('generated_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}")
        metadata_para.add_run("\nDifficulty: ").bold = True
        metadata_para.add_run(f"{quiz_data.metadata.get('difficulty', 'Unknown').title()}")
        metadata_para.add_run("\nTotal Questions: ").bold = True
        metadata_para.add_run(f"{len(quiz_data.questions)}")
        
        if quiz_data.metadata.get('topic_focus'):
            metadata_para.add_run("\nTopic Focus: ").bold = True
            metadata_para.add_run(f"{quiz_data.metadata.get('topic_focus')}")
        
        doc.add_paragraph()  # Add space
        
        # Add instructions
        instructions = doc.add_paragraph()
        instructions.add_run("Instructions: ").bold = True
        if quiz_data.include_answers:
            instructions.add_run("Choose the best answer for each question. The correct answer is indicated after the options.")
        else:
            instructions.add_run("Choose the best answer for each question.")
        
        doc.add_paragraph()  # Add space
        
        # Add questions
        for i, question_data in enumerate(quiz_data.questions, 1):
            # Question number and text
            question_para = doc.add_paragraph()
            question_para.add_run(f"Question {i}: ").bold = True
            question_para.add_run(question_data['question'])
            
            # Options
            for j, option in enumerate(question_data['options'], 1):
                option_para = doc.add_paragraph(f"   {chr(64 + j)}. {option['text']}", style='List Number')
                option_para.paragraph_format.left_indent = Inches(0.5)
            
            # Correct answer (if included)
            if quiz_data.include_answers:
                correct_option = next((opt for opt in question_data['options'] if opt['is_correct']), None)
                if correct_option:
                    correct_letter = chr(65 + next(i for i, opt in enumerate(question_data['options']) if opt['is_correct']))
                    answer_para = doc.add_paragraph()
                    answer_para.add_run("Correct Answer: ").bold = True
                    answer_para.add_run(f"{correct_letter}. {correct_option['text']}")
                    
                    # Add explanation if available
                    if question_data.get('explanation'):
                        explanation_para = doc.add_paragraph()
                        explanation_para.add_run("Explanation: ").bold = True
                        explanation_para.add_run(question_data['explanation'])
            
            # Add space between questions
            doc.add_paragraph()
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Successfully created Word document with {len(quiz_data.questions)} questions")
        return buffer
        
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        return None


def create_quiz_pdf_document(quiz_data: QuizExportData) -> Optional[io.BytesIO]:
    """
    Create a PDF document from quiz data.
    
    Args:
        quiz_data: QuizExportData containing quiz information
        
    Returns:
        BytesIO buffer containing the PDF document or None if failed
    """
    if not REPORTLAB_AVAILABLE:
        logger.error("reportlab not available for PDF document generation")
        return None
    
    try:
        # Create buffer
        buffer = io.BytesIO()
        
        # Create document
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        question_style = ParagraphStyle(
            'QuestionStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=10,
            fontName='Helvetica-Bold'
        )
        
        option_style = ParagraphStyle(
            'OptionStyle',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=5
        )
        
        answer_style = ParagraphStyle(
            'AnswerStyle',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica-Bold',
            textColor=green,
            spaceAfter=5
        )
        
        explanation_style = ParagraphStyle(
            'ExplanationStyle',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20,
            spaceAfter=15,
            textColor=blue
        )
        
        # Build content
        content = []
        
        # Title
        content.append(Paragraph(quiz_data.title, title_style))
        content.append(Spacer(1, 20))
        
        # Metadata
        metadata_text = f"""
        <b>Generated:</b> {quiz_data.metadata.get('generated_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}<br/>
        <b>Difficulty:</b> {quiz_data.metadata.get('difficulty', 'Unknown').title()}<br/>
        <b>Total Questions:</b> {len(quiz_data.questions)}
        """
        
        if quiz_data.metadata.get('topic_focus'):
            metadata_text += f"<br/><b>Topic Focus:</b> {quiz_data.metadata.get('topic_focus')}"
        
        content.append(Paragraph(metadata_text, styles['Normal']))
        content.append(Spacer(1, 20))
        
        # Instructions
        instructions_text = "<b>Instructions:</b> "
        if quiz_data.include_answers:
            instructions_text += "Choose the best answer for each question. The correct answer is indicated after the options."
        else:
            instructions_text += "Choose the best answer for each question."
        
        content.append(Paragraph(instructions_text, styles['Normal']))
        content.append(Spacer(1, 30))
        
        # Questions
        for i, question_data in enumerate(quiz_data.questions, 1):
            # Question
            question_text = f"<b>Question {i}:</b> {question_data['question']}"
            content.append(Paragraph(question_text, question_style))
            
            # Options
            for j, option in enumerate(question_data['options'], 1):
                option_text = f"{chr(64 + j)}. {option['text']}"
                content.append(Paragraph(option_text, option_style))
            
            # Correct answer (if included)
            if quiz_data.include_answers:
                correct_option = next((opt for opt in question_data['options'] if opt['is_correct']), None)
                if correct_option:
                    correct_letter = chr(65 + next(i for i, opt in enumerate(question_data['options']) if opt['is_correct']))
                    answer_text = f"<b>Correct Answer:</b> {correct_letter}. {correct_option['text']}"
                    content.append(Paragraph(answer_text, answer_style))
                    
                    # Add explanation if available
                    if question_data.get('explanation'):
                        explanation_text = f"<b>Explanation:</b> {question_data['explanation']}"
                        content.append(Paragraph(explanation_text, explanation_style))
            
            # Add space between questions
            content.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(content)
        buffer.seek(0)
        
        logger.info(f"Successfully created PDF document with {len(quiz_data.questions)} questions")
        return buffer
        
    except Exception as e:
        logger.error(f"Error creating PDF document: {e}")
        return None


def prepare_quiz_export_data(quiz_session, include_answers: bool = True) -> QuizExportData:
    """
    Prepare quiz data for export.

    Args:
        quiz_session: QuizSession object from Streamlit
        include_answers: Whether to include answers in the export

    Returns:
        QuizExportData object ready for export
    """
    questions_data = []

    # Access questions as attribute, not dictionary key
    for question in quiz_session.questions:
        question_data = {
            'question': question.question,
            'options': [{'text': opt.text, 'is_correct': opt.is_correct} for opt in question.options],
            'explanation': question.explanation,
            'topic': question.topic,
            'difficulty': question.difficulty
        }
        questions_data.append(question_data)

    # Extract difficulty from first question since QuizSession doesn't store it directly
    difficulty = questions_data[0]['difficulty'] if questions_data else 'Unknown'

    # Extract topic focus from questions (look for common topic)
    topics = [q['topic'] for q in questions_data]
    topic_focus = max(set(topics), key=topics.count) if topics else ''

    metadata = {
        'generated_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'difficulty': difficulty,
        'topic_focus': topic_focus,
        'total_questions': len(questions_data)
    }

    title = f"Quiz - {metadata['difficulty'].title()} Level"
    if metadata['topic_focus'] and metadata['topic_focus'] != 'General':
        title += f" ({metadata['topic_focus']})"

    return QuizExportData(
        title=title,
        questions=questions_data,
        metadata=metadata,
        include_answers=include_answers
    )
